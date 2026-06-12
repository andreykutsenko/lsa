#!/usr/bin/env python3
"""
CAB Questionnaire generator.

Produces a docx in the standard CAB format (Calibri 11pt, 7 sections, bullets).

Usage (library):
    from lsa.changedocs.generate_cab import build_cab
    build_cab(content_dict, "output.docx")

content_dict schema:
    {
        "ticket_id": "SP2-1167",
        "title": "Remove leading zeros for Due Diligence Letter (DL014) - MOCU Daily Notices",
        "sections": [
            {"num": 1, "name": "Client Isolation", "items": [
                {"kind": "bullet", "q": "...", "a": "..."},
                ...
            ]},
            {"num": 2, "name": "Change Classification", "items": [
                {"kind": "bullet", "q": "Is this:", "a": ""},
                {"kind": "sub",    "q": "config-only?", "a": "Yes..."},
                ...
            ]},
            ...
        ]
    }
"""

import argparse
import json

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


FONT_NAME = "Calibri"
FONT_SIZE_PT = 11


def _set_run_font(run, bold=False):
    run.bold = bold
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE_PT)


def _add_centered_bold(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    _set_run_font(r, bold=True)


def _add_empty(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)


def _add_section(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("{}. {}".format(num, text))
    _set_run_font(r, bold=True)


def _add_bullet_qa(doc, question, answer, level="bullet"):
    p = doc.add_paragraph()
    if level == "bullet":
        p.paragraph_format.left_indent = Inches(0.25)
        marker = "○ "  # ○
    else:
        p.paragraph_format.left_indent = Inches(0.5)
        marker = "■ "  # ■
    p.paragraph_format.space_after = Pt(2)

    r = p.add_run(marker)
    _set_run_font(r)

    r_q = p.add_run(question)
    _set_run_font(r_q, bold=True)

    if answer:
        r_a = p.add_run("  " + answer)
        _set_run_font(r_a)


def build_cab(content, output_path):
    """Build a CAB docx from a content dict. See module docstring for schema."""
    doc = Document()

    # Default style
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(FONT_SIZE_PT)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    _add_centered_bold(doc, "CAB Questionnaire")
    _add_empty(doc)
    _add_centered_bold(
        doc,
        "{ticket_id} – {title}".format(
            ticket_id=content["ticket_id"], title=content["title"]
        ),
    )
    _add_empty(doc)

    for section in content["sections"]:
        _add_section(doc, section["num"], section["name"])
        for item in section["items"]:
            kind = item.get("kind", "bullet")
            level = "bullet" if kind == "bullet" else "sub"
            _add_bullet_qa(doc, item["q"], item.get("a", ""), level=level)

    doc.save(output_path)
    return output_path


def _cli():
    ap = argparse.ArgumentParser(description="Generate a CAB Questionnaire docx.")
    ap.add_argument("--config", required=True, help="Path to JSON config with CAB content.")
    ap.add_argument("--output", required=True, help="Path to output docx.")
    args = ap.parse_args()

    with open(args.config, "r", encoding="utf-8") as fh:
        content = json.load(fh)

    out = build_cab(content, args.output)
    print("OK: {}".format(out))


if __name__ == "__main__":
    _cli()
