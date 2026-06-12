"""Deterministic docx rendering: CAB (from LLM content) and PTF (from context).

CAB reuses generate_cab.build_cab. PTF is filled entirely from the parallel
header plus human-owned fields (jira, hours, dates); it needs no LLM.
"""

import datetime
import os
import shutil
from copy import deepcopy

from docx import Document
from docx.shared import Pt

from .generate_cab import build_cab

_HERE = os.path.dirname(os.path.abspath(__file__))

PTF_TEMPLATE = os.path.join(_HERE, "templates", "ptf_template.docx")
QA_TEMPLATE = os.path.join(_HERE, "templates", "qa_template.docx")

# Fonts matching the reference PTF: form fields are Calibri 11pt bold; technical
# identifiers (parallel id, file names) are Segoe UI 10.5pt regular. The template's
# value cells are empty, so without these the text would inherit Times New Roman.
_FORM_FONT = ("Calibri", Pt(11))
_ID_FONT = ("Segoe UI", Pt(10.5))

DEFAULT_QA_ITEMS = [
    "Ensure that all the script updates in the assessment have been completed.",
]


def render_cab(cab_content, out_path):
    """cab_content: dict with ticket_id, title, sections (generate_cab schema)."""
    build_cab(cab_content, out_path)
    return out_path


def _set_cell(cell, text, font_name=None, font_size=None, bold=None):
    """Replace a cell's text with an explicit/inherited font.

    Explicit font_name/font_size/bold win; otherwise the first run's formatting is
    reused; otherwise a Calibri 11pt default is applied so empty template cells do
    not fall back to the document's Times New Roman Normal style.
    """
    paragraphs = list(cell.paragraphs)
    for p in paragraphs[1:]:
        p._p.getparent().remove(p._p)
    p = cell.paragraphs[0]
    captured = None
    if p.runs:
        r0 = p.runs[0]
        captured = dict(name=r0.font.name, size=r0.font.size, bold=r0.bold, italic=r0.italic)
        for r in list(p.runs):
            r._r.getparent().remove(r._r)
    r = p.add_run(text)
    r.font.name = font_name or (captured and captured["name"]) or "Calibri"
    r.font.size = font_size or (captured and captured["size"]) or Pt(11)
    r.bold = bold if bold is not None else (captured["bold"] if captured else None)
    r.italic = captured["italic"] if captured else None


def render_ptf(context, out_path, jira="", hours="", live_date="", date=None,
               programmer="", template=PTF_TEMPLATE):
    """Fill the PTF template from the parallel header + human-owned fields."""
    if date is None:
        date = datetime.date.today().strftime("%m/%d/%Y")
    shutil.copy(template, out_path)
    d = Document(out_path)
    t0, t1, t2 = d.tables[0], d.tables[1], d.tables[2]

    form_name, form_size = _FORM_FONT
    id_name, id_size = _ID_FONT
    _set_cell(t0.rows[0].cells[6], " " + jira, font_name=form_name, font_size=form_size, bold=True)
    _set_cell(t0.rows[1].cells[6], " " + hours, font_name=form_name, font_size=form_size, bold=True)
    _set_cell(t0.rows[2].cells[1], " " + programmer, font_name=form_name, font_size=form_size, bold=True)
    _set_cell(t0.rows[2].cells[3], " " + date, font_name=form_name, font_size=form_size, bold=True)
    _set_cell(t0.rows[2].cells[6], " " + (live_date or date), font_name=form_name, font_size=form_size, bold=True)
    _set_cell(t0.rows[3].cells[1], " " + context.get("parallel_id", ""), font_name=id_name, font_size=id_size, bold=False)
    _set_cell(t1.rows[1].cells[0], context.get("description", ""), font_name=form_name, font_size=form_size, bold=False)

    files = context.get("files", [])
    for i, fname in enumerate(files):
        if 1 + i < len(t2.rows):
            _set_cell(t2.rows[1 + i].cells[0], fname, font_name=id_name, font_size=id_size, bold=False)
    _set_cell(t2.rows[len(t2.rows) - 1].cells[3],
              "Total # File(s) Transferred: {}".format(len(files)),
              font_name=form_name, font_size=form_size, bold=True)

    d.save(out_path)
    return out_path


def render_qa(context, out_path, job_number="", date=None, l_items=None,
              programmer="", template=QA_TEMPLATE):
    """Fill the QA checklist template from the header + L test-case items.

    Deterministic — no API call. l_items defaults to a single generic item;
    pass a list to add L1..Ln (the template's L1 row is cloned per item).
    """
    if date is None:
        date = datetime.date.today().strftime("%m/%d/%Y")
    if not l_items:
        l_items = list(DEFAULT_QA_ITEMS)
    if not job_number:
        desc = context.get("description", "")
        job_number = desc.split()[0] if desc else context.get("parallel_id", "")

    shutil.copy(template, out_path)
    d = Document(out_path)
    t0, t1 = d.tables[0], d.tables[1]

    _set_cell(t0.rows[0].cells[2], programmer)
    _set_cell(t0.rows[0].cells[4], job_number)
    _set_cell(t0.rows[0].cells[6], date)
    _set_cell(t0.rows[1].cells[2], context.get("description", ""))

    # Locate the L1 row and the Peer/QA sign-off row that follows the L block.
    l1_idx = peer_idx = None
    for i, row in enumerate(t1.rows):
        label = row.cells[0].text.strip()
        if label == "L1" and l1_idx is None:
            l1_idx = i
        if label.startswith("Peer / QA Sign Off"):
            peer_idx = i
            break
    if l1_idx is None or peer_idx is None:
        raise ValueError("QA template missing L1 / Peer sign-off rows.")

    l1_row = t1.rows[l1_idx]
    peer_row = t1.rows[peer_idx]
    for _ in range(len(l_items) - 1):
        peer_row._tr.addprevious(deepcopy(l1_row._tr))

    for offset, text in enumerate(l_items):
        row = t1.rows[l1_idx + offset]
        _set_cell(row.cells[0], "L{}".format(offset + 1))
        _set_cell(row.cells[1], text)

    d.save(out_path)
    return out_path
