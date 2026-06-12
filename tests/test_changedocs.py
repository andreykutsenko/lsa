"""Offline / deterministic tests for the ported Change Docs engine and routes.

No live ssh and no live Claude API calls — context.fetch and draft.draft_cab
are monkeypatched. Route coroutines are awaited directly following the existing
tests/test_web_server.py convention.
"""

import inspect

import pytest
from docx import Document

from lsa.changedocs import context as ctx
from lsa.changedocs import draft as drafter
from lsa.changedocs import render as renderer
from lsa.web import server


# --- AC3: whitelist + size caps --------------------------------------------

def test_collect_diffs_whitelist_and_caps():
    big_code = "x" * (ctx.MAX_FILE_DIFF_BYTES + 100)
    filler = "y" * (ctx.MAX_FILE_DIFF_BYTES)
    pairs = [
        ("good.py", "real diff content"),
        ("data.csv", "should be skipped (non-code)"),
        ("empty.sh", "   \n  "),
        ("huge.dfa", big_code),
        # Push total over the cap: several near-max whitelisted files.
        ("a.pl", filler),
        ("b.pl", filler),
        ("c.pl", filler),
        ("d.pl", filler),
    ]
    diffs, skipped = ctx._collect_diffs(pairs)

    by_file = {d["file"]: d for d in diffs}
    assert "good.py" in by_file
    assert by_file["good.py"]["truncated"] is False
    assert "huge.dfa" in by_file
    assert by_file["huge.dfa"]["truncated"] is True
    assert len(by_file["huge.dfa"]["diff"]) == ctx.MAX_FILE_DIFF_BYTES

    assert any("data.csv" in s and "non-code" in s for s in skipped)
    assert any("empty.sh" in s and "empty" in s for s in skipped)
    assert any("total diff size cap" in s for s in skipped)

    total = sum(len(d["diff"]) for d in diffs)
    assert total <= ctx.MAX_TOTAL_DIFF_BYTES


def test_ported_constants_unchanged():
    assert ctx.MAX_TOTAL_DIFF_BYTES == 200_000
    assert ctx.MAX_FILE_DIFF_BYTES == 60_000
    assert ctx.CODE_EXTENSIONS == {
        "dfa", "sh", "pl", "py", "procs", "control", "ovl", "ins", "lis", "300",
    }


# --- AC4: header parsing ----------------------------------------------------

def test_parse_header_extracts_fields():
    report = (
        "Some banner\n"
        "Description: SP1-2045 fix leading zeros\n"
        "Parallel ID: 20260520083853\n"
        "Files:\n"
        "  1) procs/crcums1.procs\n"
        "  2) docdef/crcu.dfa\n"
        "* end of files\n"
        "  3) ignored.after.star\n"
    )
    header = ctx.parse_header(report)
    assert header["description"] == "SP1-2045 fix leading zeros"
    assert header["parallel_id"] == "20260520083853"
    assert header["files"] == ["procs/crcums1.procs", "docdef/crcu.dfa"]


# --- AC9 (part): PRID validation -------------------------------------------

def test_fetch_rejects_short_prid():
    with pytest.raises(ctx.ContextError):
        ctx.fetch("123")


# --- AC5: dry-run makes no API call, returns a cost estimate ----------------

def test_dry_run_returns_estimate_no_api():
    prompt = "Parallel ID: 1\nDescription: x\nDiffs:\n### a.py\nreal"
    est = drafter.dry_run(prompt, model="sonnet")
    assert est["model_id"] == drafter.MODELS["sonnet"]
    assert est["estimated_input_tokens"] > 0
    assert est["max_output_tokens"] == drafter.MAX_OUTPUT_TOKENS
    assert isinstance(est["estimated_cost_usd"], float)
    assert est["estimated_cost_usd"] > 0


@pytest.mark.anyio
async def test_preview_route_no_api_no_usage_log(monkeypatch, tmp_path):
    sample = {
        "parallel_id": "20260520083853",
        "description": "SP1-2045 fix",
        "files": ["a.py"],
        "diffs": [{"file": "a.py", "diff": "real diff", "truncated": False}],
        "skipped": ["data.csv (non-code extension)"],
    }
    monkeypatch.setattr(ctx, "fetch", lambda prid, remote=None: sample)

    def _fail_draft(*a, **k):
        raise AssertionError("draft_cab must not be called on preview")

    monkeypatch.setattr(drafter, "draft_cab", _fail_draft)

    usage_log = tmp_path / "usage.log"
    monkeypatch.setattr(drafter, "USAGE_LOG_PATH", str(usage_log))
    monkeypatch.setattr("lsa.config.load_user_config", lambda: {})

    req = server.ChangeDocsPreviewRequest(prid="20260520083853", model="sonnet")
    resp = await server.changedocs_preview(req)

    assert resp["no_api_call"] is True
    assert resp["model_id"] == drafter.MODELS["sonnet"]
    assert resp["estimated_input_tokens"] > 0
    assert resp["max_output_tokens"] == drafter.MAX_OUTPUT_TOKENS
    assert resp["estimated_cost_usd"] > 0
    assert not usage_log.exists()


# --- AC6: generate writes requested files, returns download refs ------------

def _fake_cab_content():
    return {
        "ticket_id": "SP1-2045",
        "title": "Fix leading zeros",
        "sections": [
            {"num": 1, "name": "Client Isolation", "items": [
                {"kind": "bullet", "q": "Q?", "a": "A."},
            ]},
        ],
    }


def _fake_context():
    return {
        "parallel_id": "20260520083853",
        "description": "SP1-2045 fix",
        "files": ["procs/a.procs", "docdef/b.dfa"],
        "diffs": [{"file": "procs/a.procs", "diff": "real", "truncated": False}],
        "skipped": [],
    }


def _patch_generate(monkeypatch, tmp_path):
    monkeypatch.setattr(ctx, "fetch", lambda prid, remote=None: _fake_context())
    monkeypatch.setattr(
        drafter,
        "draft_cab",
        lambda context, model="sonnet", **kwargs: _fake_cab_content(),
    )
    monkeypatch.setattr(
        server, "_changedocs_out_root", lambda cfg: tmp_path / "changedocs"
    )
    monkeypatch.setattr("lsa.config.load_user_config", lambda: {})


@pytest.mark.anyio
async def test_generate_cab_only(monkeypatch, tmp_path):
    _patch_generate(monkeypatch, tmp_path)
    req = server.ChangeDocsGenerateRequest(prid="20260520083853", ptf=False, qa=False)
    resp = await server.changedocs_generate(req)

    kinds = [f["kind"] for f in resp["files"]]
    assert kinds == ["cab"]
    out_dir = tmp_path / "changedocs" / resp["ticket_id"]
    assert (out_dir / f"{resp['ticket_id']}_CAB_Questionnaire.docx").exists()
    assert not (out_dir / f"{resp['ticket_id']}_PTF.docx").exists()
    for f in resp["files"]:
        assert f["download"].startswith("/api/changedocs/download?")


@pytest.mark.anyio
async def test_generate_cab_ptf_qa(monkeypatch, tmp_path):
    _patch_generate(monkeypatch, tmp_path)
    req = server.ChangeDocsGenerateRequest(
        prid="20260520083853", ptf=True, qa=True,
        jira="SP1-2045", hours="8hrs", live_date="06/10/2026",
    )
    resp = await server.changedocs_generate(req)

    kinds = sorted(f["kind"] for f in resp["files"])
    assert kinds == ["cab", "ptf", "qa"]
    out_dir = tmp_path / "changedocs" / resp["ticket_id"]
    for suffix in ("_CAB_Questionnaire.docx", "_PTF.docx", "_QA_Checklist.docx"):
        assert (out_dir / f"{resp['ticket_id']}{suffix}").exists()


# --- AC7: deterministic PTF / QA rendering ----------------------------------

def test_render_ptf_fills_fields(tmp_path):
    out = tmp_path / "ptf.docx"
    renderer.render_ptf(_fake_context(), str(out), jira="SP1-2045",
                        hours="8hrs", live_date="06/10/2026")
    assert out.exists()
    doc = Document(str(out))
    text = "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    assert "SP1-2045" in text
    assert "8hrs" in text
    assert "20260520083853" in text
    assert "Total # File(s) Transferred: 2" in text


def test_render_ptf_uses_reference_fonts(tmp_path):
    """Filled cells must match the reference PTF fonts, not inherit Times New Roman."""
    out = tmp_path / "ptf.docx"
    renderer.render_ptf(_fake_context(), str(out), jira="SP1-2045", hours="8hrs")
    doc = Document(str(out))
    t0, _, t2 = doc.tables

    def font_of(cell):
        run = cell.paragraphs[0].runs[0]
        return run.font.name, (run.font.size.pt if run.font.size else None), run.bold

    # Form fields: Calibri 11 bold.
    assert font_of(t0.rows[0].cells[6]) == ("Calibri", 11.0, True)   # jira
    assert font_of(t0.rows[2].cells[1]) == ("Calibri", 11.0, True)   # programmer
    # Technical identifiers: Segoe UI 10.5, not bold.
    assert font_of(t0.rows[3].cells[1]) == ("Segoe UI", 10.5, False)  # parallel id
    assert font_of(t2.rows[1].cells[0]) == ("Segoe UI", 10.5, False)  # file name
    # Never the document default Times New Roman.
    for cell in (t0.rows[0].cells[6], t0.rows[3].cells[1]):
        assert cell.paragraphs[0].runs[0].font.name != "Times New Roman"


def test_render_qa_clones_l_rows(tmp_path):
    out = tmp_path / "qa.docx"
    items = ["First check", "Second check", "Third check"]
    renderer.render_qa(_fake_context(), str(out), job_number="JOB-1", l_items=items)
    assert out.exists()
    doc = Document(str(out))
    t1 = doc.tables[1]
    labels = {row.cells[0].text.strip() for row in t1.rows}
    assert {"L1", "L2", "L3"} <= labels
    cell_text = "\n".join(
        cell.text for row in t1.rows for cell in row.cells
    )
    for item in items:
        assert item in cell_text


# --- AC8: download route is path-safe ---------------------------------------

@pytest.mark.anyio
async def test_download_streams_and_rejects_traversal(monkeypatch, tmp_path):
    out_root = tmp_path / "changedocs"
    ticket_dir = out_root / "SP1-2045"
    ticket_dir.mkdir(parents=True)
    target = ticket_dir / "SP1-2045_CAB_Questionnaire.docx"
    target.write_bytes(b"PK\x03\x04 fake docx")

    monkeypatch.setattr(server, "_changedocs_out_root", lambda cfg: out_root)
    monkeypatch.setattr("lsa.config.load_user_config", lambda: {})

    resp = await server.changedocs_download(
        ticket_id="SP1-2045", name="SP1-2045_CAB_Questionnaire.docx"
    )
    assert resp.path == str(target.resolve())
    assert resp.media_type == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    # Re-download succeeds.
    resp2 = await server.changedocs_download(
        ticket_id="SP1-2045", name="SP1-2045_CAB_Questionnaire.docx"
    )
    assert resp2.path == str(target.resolve())

    from fastapi import HTTPException

    with pytest.raises(HTTPException) as missing:
        await server.changedocs_download(ticket_id="SP1-2045", name="nope.docx")
    assert missing.value.status_code == 404

    with pytest.raises(HTTPException) as traversal:
        await server.changedocs_download(ticket_id="SP1-2045", name="../../etc/passwd")
    assert traversal.value.status_code in (403, 404)


# --- AC9: guardrails intact in code -----------------------------------------

def test_draft_guardrails_present_in_source():
    src = inspect.getsource(drafter.draft_cab)
    assert 'os.environ.get("ANTHROPIC_API_KEY")' in src
    assert "tools" not in src  # no tools handed to the model
    assert "cache_control" in src
    assert "ephemeral" in src
    assert "max_tokens=MAX_OUTPUT_TOKENS" in src
    # Exactly one repair retry on malformed JSON: two invocation sites
    # (resp = _call() then resp = _call("...")), excluding the def.
    assert src.count("resp = _call(") == 2
    assert "_log_usage(" in src


def test_default_model_and_output_cap():
    assert drafter.DEFAULT_MODEL == "sonnet"
    assert drafter.MAX_OUTPUT_TOKENS == 4000
    assert drafter.MODELS["sonnet"] == "claude-sonnet-4-6"


# --- API key file: status / save / delete -----------------------------------

@pytest.mark.anyio
async def test_key_status_save_and_delete(monkeypatch, tmp_path):
    import os
    key_file = tmp_path / "anthropic_key"
    monkeypatch.setattr(server, "_ANTHROPIC_KEY_FILE", key_file)
    monkeypatch.delenv("ANTHROPIC_API_KEY", raising=False)

    # Unconfigured to begin with.
    status = await server.changedocs_key_status()
    assert status == {"configured": False, "source": None, "masked": None}

    # Save a valid key -> persisted, reported configured, masked (not revealed).
    full_key = "sk-ant-api03-" + "A" * 40
    saved = await server.changedocs_key_save(server.ApiKeyRequest(api_key=full_key))
    assert saved["configured"] is True
    assert saved["source"] == "saved"
    assert full_key not in saved["masked"]
    assert key_file.read_text() == full_key
    assert (key_file.stat().st_mode & 0o777) == 0o600

    # Delete -> back to unconfigured, file gone.
    after = await server.changedocs_key_delete()
    assert after["configured"] is False
    assert not key_file.exists()


@pytest.mark.anyio
async def test_key_save_rejects_non_anthropic(monkeypatch, tmp_path):
    from fastapi import HTTPException
    key_file = tmp_path / "anthropic_key"
    monkeypatch.setattr(server, "_ANTHROPIC_KEY_FILE", key_file)

    with pytest.raises(HTTPException) as bad:
        await server.changedocs_key_save(server.ApiKeyRequest(api_key="not-a-key"))
    assert bad.value.status_code == 400
    assert not key_file.exists()


def test_detail_concise_suffix_and_estimate():
    assert drafter._style_suffix("detailed") == ""
    assert drafter.CONCISE_STYLE_DIRECTIVE in drafter._style_suffix("concise")
    detailed = drafter.dry_run("base prompt", model="sonnet", detail="detailed")
    concise = drafter.dry_run("base prompt", model="sonnet", detail="concise")
    assert concise["context_tokens"] > detailed["context_tokens"]


def test_detail_defaults_to_concise():
    assert server.ChangeDocsGenerateRequest(prid="20260520083853").detail == "concise"
    assert server.ChangeDocsPreviewRequest(prid="20260520083853").detail == "concise"


def test_extra_context_block_and_estimate():
    assert drafter._extra_block("") == ""
    assert drafter._extra_block("   ") == ""
    assert "operator note" in drafter._extra_block("operator note")
    base = drafter.dry_run("base prompt", model="sonnet", extra_context="")
    withctx = drafter.dry_run("base prompt", model="sonnet", extra_context="extra note here")
    assert withctx["context_tokens"] > base["context_tokens"]


def test_generate_uses_saved_key_when_field_blank(monkeypatch, tmp_path):
    """generate resolves api_key: request field -> saved file -> env."""
    key_file = tmp_path / "anthropic_key"
    key_file.write_text("sk-ant-api03-" + "B" * 40)
    monkeypatch.setattr(server, "_ANTHROPIC_KEY_FILE", key_file)
    assert server._load_saved_api_key() == "sk-ant-api03-" + "B" * 40


@pytest.mark.anyio
async def test_download_rejects_sibling_prefix_dir(monkeypatch, tmp_path):
    out_root = tmp_path / "changedocs"
    out_root.mkdir()
    evil_dir = tmp_path / "changedocs_evil"
    evil_dir.mkdir()
    (evil_dir / "x.docx").write_bytes(b"PK\x03\x04 fake docx")

    monkeypatch.setattr(server, "_changedocs_out_root", lambda cfg: out_root)
    monkeypatch.setattr("lsa.config.load_user_config", lambda: {})

    from fastapi import HTTPException

    with pytest.raises(HTTPException) as exc:
        await server.changedocs_download(ticket_id="..", name="changedocs_evil/x.docx")
    assert exc.value.status_code == 403


@pytest.mark.anyio
async def test_generate_sanitizes_ticket_id(monkeypatch, tmp_path):
    _patch_generate(monkeypatch, tmp_path)
    req = server.ChangeDocsGenerateRequest(
        prid="20260520083853", ticket_id="../../evil", ptf=False, qa=False
    )
    resp = await server.changedocs_generate(req)

    assert "/" not in resp["ticket_id"]
    assert "\\" not in resp["ticket_id"]
    out_root = (tmp_path / "changedocs").resolve()
    from pathlib import Path as _Path

    assert _Path(resp["out_dir"]).resolve().is_relative_to(out_root)


# --- medium-fixes-v1: error handling and usage accounting ---------------------

def test_fetch_converts_ssh_timeout_to_context_error(monkeypatch):
    import subprocess

    def _timeout(*args, **kwargs):
        raise subprocess.TimeoutExpired(cmd="ssh", timeout=300)

    monkeypatch.setattr(ctx.subprocess, "run", _timeout)
    with pytest.raises(ctx.ContextError, match="timed out"):
        ctx.fetch("20260520083853")


class _FakeResp:
    def __init__(self, text, in_tok=0, out_tok=0):
        from types import SimpleNamespace

        self.content = [SimpleNamespace(type="text", text=text)]
        self.usage = SimpleNamespace(input_tokens=in_tok, output_tokens=out_tok)


def _install_fake_anthropic(monkeypatch, responses):
    import anthropic

    class _Messages:
        def create(self, **kwargs):
            item = responses.pop(0)
            if isinstance(item, Exception):
                raise item
            return item

    class _Client:
        def __init__(self, api_key=None):
            self.messages = _Messages()

    monkeypatch.setattr(anthropic, "Anthropic", _Client)


def _read_usage_line(path):
    line = path.read_text().strip().splitlines()[-1]
    fields = dict(f.split("=", 1) for f in line.split("\t")[1:] if "=" in f)
    return int(fields["in"]), int(fields["out"])


def test_draft_cab_converts_api_error(monkeypatch, tmp_path):
    import anthropic
    import httpx

    err = anthropic.APIConnectionError(
        request=httpx.Request("POST", "https://api.anthropic.com")
    )
    _install_fake_anthropic(monkeypatch, [err])
    with pytest.raises(drafter.DraftError, match="API call failed"):
        drafter.draft_cab(
            _fake_context(), api_key="sk-test",
            usage_log_path=str(tmp_path / "usage.log"),
        )


def test_draft_cab_raises_draft_error_on_double_malformed_json(monkeypatch, tmp_path):
    _install_fake_anthropic(monkeypatch, [
        _FakeResp("not json", 100, 10),
        _FakeResp("still not json", 120, 20),
    ])
    log = tmp_path / "usage.log"
    with pytest.raises(drafter.DraftError, match="malformed JSON"):
        drafter.draft_cab(_fake_context(), api_key="sk-test", usage_log_path=str(log))
    assert _read_usage_line(log) == (220, 30)


def test_draft_cab_accumulates_usage_across_repair_retry(monkeypatch, tmp_path):
    _install_fake_anthropic(monkeypatch, [
        _FakeResp("not json", 100, 10),
        _FakeResp('{"sections": [{"num": 1}]}', 150, 40),
    ])
    log = tmp_path / "usage.log"
    content = drafter.draft_cab(
        _fake_context(), api_key="sk-test", usage_log_path=str(log)
    )
    assert content["sections"] == [{"num": 1}]
    assert _read_usage_line(log) == (250, 50)


# --- changedocs-fixes-v2 ------------------------------------------------------

def test_remote_cleanup_failure_is_non_fatal(monkeypatch):
    """The remote command must tolerate rm failing (NFS lock files)."""
    from types import SimpleNamespace

    captured = {}

    def _fake_run(cmd, **kwargs):
        captured["remote_cmd"] = cmd[2]
        return SimpleNamespace(returncode=0, stdout="__NOFOLDER__", stderr="")

    monkeypatch.setattr(ctx.subprocess, "run", _fake_run)
    with pytest.raises(ctx.ContextError, match="No parallel folder"):
        ctx.fetch("20260520083853")
    assert 'rm -rf "$F" >/dev/null 2>&1 || true' in captured["remote_cmd"]


@pytest.mark.anyio
async def test_generate_without_cab_renders_ptf_qa_only(monkeypatch, tmp_path):
    _patch_generate(monkeypatch, tmp_path)

    def _no_draft(*args, **kwargs):
        raise AssertionError("draft_cab must not be called when cab=False")

    monkeypatch.setattr(drafter, "draft_cab", _no_draft)
    req = server.ChangeDocsGenerateRequest(
        prid="20260520083853", cab=False, ptf=True, qa=True, ticket_id="SP1-77"
    )
    resp = await server.changedocs_generate(req)

    kinds = sorted(f["kind"] for f in resp["files"])
    assert kinds == ["ptf", "qa"]
    assert resp["ticket_id"] == "SP1-77"
    out_dir = tmp_path / "changedocs" / "SP1-77"
    assert (out_dir / "SP1-77_PTF.docx").exists()
    assert not (out_dir / "SP1-77_CAB_Questionnaire.docx").exists()
    assert "out_dir_win" in resp


@pytest.mark.anyio
async def test_generate_requires_at_least_one_document(monkeypatch, tmp_path):
    _patch_generate(monkeypatch, tmp_path)
    from fastapi import HTTPException

    req = server.ChangeDocsGenerateRequest(
        prid="20260520083853", cab=False, ptf=False, qa=False
    )
    with pytest.raises(HTTPException) as exc:
        await server.changedocs_generate(req)
    assert exc.value.status_code == 400


def _fake_openai_response(payload, status=200):
    import httpx

    class _Resp:
        status_code = status
        text = "error body"

        def raise_for_status(self):
            if status >= 400:
                raise httpx.HTTPStatusError(
                    "boom", request=httpx.Request("POST", "https://x"),
                    response=httpx.Response(status, request=httpx.Request("POST", "https://x"),
                                            text="error body"),
                )

        def json(self):
            return payload

    return _Resp()


def test_draft_cab_openai_success_and_usage(monkeypatch, tmp_path):
    import httpx

    calls = []

    def _fake_post(url, **kwargs):
        calls.append((url, kwargs))
        return _fake_openai_response({
            "choices": [{"message": {"content": '{"sections": [{"num": 1}]}'}}],
            "usage": {"prompt_tokens": 120, "completion_tokens": 30},
        })

    monkeypatch.setattr(httpx, "post", _fake_post)
    log = tmp_path / "usage.log"
    content = drafter.draft_cab(
        _fake_context(), provider="openai", model="gpt-4o-mini",
        api_key="sk-test", usage_log_path=str(log),
    )
    assert content["sections"] == [{"num": 1}]
    assert calls[0][0] == "https://api.openai.com/v1/chat/completions"
    assert calls[0][1]["json"]["model"] == "gpt-4o-mini"
    line = log.read_text()
    assert "model=openai:gpt-4o-mini" in line
    assert "in=120" in line and "out=30" in line


def test_draft_cab_openai_http_error_raises_draft_error(monkeypatch, tmp_path):
    import httpx

    monkeypatch.setattr(httpx, "post", lambda url, **kw: _fake_openai_response({}, status=401))
    with pytest.raises(drafter.DraftError, match="HTTP 401"):
        drafter.draft_cab(_fake_context(), provider="openai", api_key="sk-bad",
                          usage_log_path=str(tmp_path / "usage.log"))


def test_draft_cab_openai_network_error_raises_draft_error(monkeypatch, tmp_path):
    import httpx

    def _boom(url, **kw):
        raise httpx.ConnectError("connection refused")

    monkeypatch.setattr(httpx, "post", _boom)
    with pytest.raises(drafter.DraftError, match="API call failed"):
        drafter.draft_cab(_fake_context(), provider="openai", api_key="sk-x",
                          usage_log_path=str(tmp_path / "usage.log"))


def test_dry_run_openai_unknown_model_has_no_cost_estimate():
    est = drafter.dry_run("prompt", provider="openai", model="some-local-model")
    assert est["model_id"] == "some-local-model"
    assert est["estimated_cost_usd"] is None


def test_resolve_model_defaults():
    assert drafter._resolve_model("anthropic", "sonnet") == "claude-sonnet-4-6"
    assert drafter._resolve_model("anthropic", "opus") == "claude-opus-4-8"
    assert drafter._resolve_model("openai", "") == drafter.DEFAULT_OPENAI_MODEL
    assert drafter._resolve_model("openai", "sonnet") == drafter.DEFAULT_OPENAI_MODEL
    assert drafter._resolve_model("openai", "gpt-4o") == "gpt-4o"


@pytest.mark.anyio
async def test_openai_key_endpoints_separate_from_anthropic(monkeypatch, tmp_path):
    monkeypatch.setattr(server, "_ANTHROPIC_KEY_FILE", tmp_path / "anthropic_key")
    monkeypatch.setattr(server, "_OPENAI_KEY_FILE", tmp_path / "openai_key")
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)
    monkeypatch.delenv("ANTHROPIC_API_KEY", raising=False)

    saved = await server.changedocs_key_save(
        server.ApiKeyRequest(api_key="sk-or-v1-" + "B" * 30, provider="openai")
    )
    assert saved["configured"] is True
    assert (tmp_path / "openai_key").exists()
    assert not (tmp_path / "anthropic_key").exists()

    anth = await server.changedocs_key_status(provider="anthropic")
    assert anth["configured"] is False

    removed = await server.changedocs_key_delete(provider="openai")
    assert removed["configured"] is False
    assert not (tmp_path / "openai_key").exists()

    from fastapi import HTTPException
    with pytest.raises(HTTPException):
        await server.changedocs_key_status(provider="bogus")


def test_out_root_resolution_order(monkeypatch, tmp_path):
    setting = tmp_path / "changedocs_outdir"
    monkeypatch.setattr(server, "_CHANGEDOCS_OUTDIR_FILE", setting)

    # Default: ~/.lsa/changedocs (never under snaproot)
    root = server._changedocs_out_root({"snaproot": str(tmp_path / "snaps")})
    assert root == server.Path.home() / ".lsa" / "changedocs"

    # Config out_root wins over default
    root = server._changedocs_out_root({"changedocs": {"out_root": str(tmp_path / "cfg")}})
    assert root == tmp_path / "cfg"

    # Saved UI setting wins over config
    setting.write_text(str(tmp_path / "saved"))
    root = server._changedocs_out_root({"changedocs": {"out_root": str(tmp_path / "cfg")}})
    assert root == tmp_path / "saved"


@pytest.mark.anyio
async def test_outdir_endpoints_save_and_validate(monkeypatch, tmp_path):
    setting = tmp_path / "changedocs_outdir"
    monkeypatch.setattr(server, "_CHANGEDOCS_OUTDIR_FILE", setting)
    monkeypatch.setattr("lsa.config.load_user_config", lambda: {})

    from fastapi import HTTPException
    with pytest.raises(HTTPException) as rel:
        await server.changedocs_outdir_set(server.OutDirRequest(out_dir="relative/path"))
    assert rel.value.status_code == 400

    target = tmp_path / "win_docs"
    resp = await server.changedocs_outdir_set(server.OutDirRequest(out_dir=str(target)))
    assert resp["out_dir"] == str(target)
    assert target.is_dir()
    assert setting.read_text() == str(target)

    current = await server.changedocs_outdir_get()
    assert current["out_dir"] == str(target)


@pytest.mark.anyio
async def test_snapshots_listing_skips_non_snapshot_dirs(monkeypatch, tmp_path):
    snap = tmp_path / "rhs_snapshot_a"
    (snap / "procs").mkdir(parents=True)
    stray = tmp_path / "changedocs"
    stray.mkdir()
    (stray / "SP1-1_PTF.docx").write_bytes(b"PK")

    monkeypatch.setattr(server, "_snaproot", tmp_path)
    monkeypatch.setattr(server, "_snapshot_path", None)

    listed = await server.list_snapshots()
    names = [s["name"] for s in listed]
    assert "rhs_snapshot_a" in names
    assert "changedocs" not in names
